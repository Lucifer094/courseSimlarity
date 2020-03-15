# -*- coding:utf-8 -*-
############# 文本格式
# 文本格式转化为txt格式
# 支持doc，docx，pdf
# 输入：
#   transform_to_txt(sourcePath, store_path)
# 参数：
#   sourcePath：源文件所在目录
#   store_path：存储路径所在目录
# 输出：
#   格式转化完成后存储文件
############# by k 2019/9/12
import os
import re
from win32com import client as doc
import docx
import codecs
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter


def file_type(file_path):
    # 判断是否为DOC文档
    flag = r".doc$"
    result = re.findall(flag, file_path)
    if result:
        return 'doc'
    # 判断是否为DOCX文档
    flag = r".docx$"
    result = re.findall(flag, file_path)
    if result:
        return 'docx'
    # 判断是否为PDF文件
    flag = r".pdf$"
    result = re.findall(flag, file_path)
    if result:
        return 'pdf'
    return 'Can not find the type of:' + file_path


# DOC转化为DOCX
def doc_to_docx(source_file_path, store_file_path):
    word = doc.Dispatch('Word.Application')
    doc_file = word.Documents.Open(FileName=source_file_path, Encoding='gb18030')
    doc_file.SaveAs(store_file_path, 12, False, "", True, "", False, False, False, False)
    doc_file.Close()
    word.Quit()


# DOCX转化为TXT
def docx_to_txt(source_file_path, test_file_path):
    docx_file = docx.Document(source_file_path)
    txt_file = codecs.open(test_file_path, 'w', encoding='utf-8')
    # 输出每一段的内容
    for para in docx_file.paragraphs:
        txt_file.write(para.text)
    txt_file.close()


# PDF转化为TXT
def pdf_to_txt(source_file_path, store_file_path):
    # 获取文档对象
    pdf0 = open(source_file_path, 'rb')
    txt_file = codecs.open(store_file_path, 'w', encoding='utf-8')
    # 创建一个与文档关联的解释器
    parser = PDFParser(pdf0)
    # 创建一个PDF文档对象
    doc = PDFDocument()
    # 连接两者
    parser.set_document(doc)
    doc.set_parser(parser)
    # 文档初始化
    doc.initialize('')
    # 创建PDF资源管理器
    resources = PDFResourceManager()
    # 创建参数分析器
    laparam = LAParams()
    # 创建一个聚合器，并接收资源管理器，参数分析器作为参数
    device = PDFPageAggregator(resources, laparams=laparam)
    # 创建一个页面解释器
    interpreter = PDFPageInterpreter(resources, device)
    # 使用文档对象获取页面的集合
    for page in doc.get_pages():
        # 使用页面解释器读取页面
        interpreter.process_page(page)
        # 使用聚合器读取页面页面内容
        layout = device.get_result()
        for out in layout:
            # 因为文档中不只有text文本
            if hasattr(out, 'get_text'):
                if out.get_text() != '\n':
                    # txt_file.write(out.get_text())
                    text = re.sub(' +', '', out.get_text())
                    if text != '\n':
                        txt_file.write(text)
    txt_file.close()


# 将文档转化为txt类型文本
def transform_to_txt(source_path, store_path):
    # 原始文档列表
    file_list = os.listdir(source_path)
    # 每一文档进行格式转化（TXT）
    for file_name in file_list:
        # 文件路径
        source_file_path = source_path + '\\' + file_name
        type_name = file_type(source_file_path)
        # doc转txt文件
        if type_name == 'doc':
            # doc转docx文件
            file_name_new = re.sub(r'doc', 'docx', file_name)
            store_file_path = source_path + '\\' + file_name_new
            doc_to_docx(source_file_path, store_file_path)
            os.remove(source_file_path)
            # docx转txt文件
            source_file_path = store_file_path
            file_name = re.sub(r'docx', 'txt', file_name_new)
            store_file_path = store_path + r'/' + file_name
            docx_to_txt(source_file_path, store_file_path)
        # docx转txt文件
        elif type_name == 'docx':
            file_name = re.sub(r'docx', 'txt', file_name)
            store_file_path = store_path + r'/' + file_name
            docx_to_txt(source_file_path, store_file_path)
        # PDF转txt文件
        elif type_name == 'pdf':
            file_name = re.sub(r'pdf', 'txt', file_name)
            store_file_path = store_path + r'/' + file_name
            pdf_to_txt(source_file_path, store_file_path)
        else:
            print(type_name)
    print("文档转化完成！")
